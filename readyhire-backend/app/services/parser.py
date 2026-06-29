import docx
from io import BytesIO

class ParsingError(ValueError):
    """Custom exception raised when document parsing fails."""
    pass

def parse_pdf(file_bytes: bytes) -> str:
    """
    Parses a PDF file in-memory.
    Attempts to use PyMuPDF (fitz) first. If it fails to import or load due to
    system security/DLL policies, it falls back to the pure-Python pypdf library.
    """
    if not file_bytes or len(file_bytes) == 0:
        raise ParsingError("The uploaded PDF file is empty.")
        
    # Attempt PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            if doc.page_count == 0:
                raise ParsingError("The PDF document has no pages.")
            
            text_parts = []
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)
            
            text = "\n".join(text_parts).strip()
            if text:
                return text
    except (ImportError, ModuleNotFoundError, Exception) as e:
        # If PyMuPDF fails due to DLL loading policy or other errors, fall back to pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(BytesIO(file_bytes))
            if len(reader.pages) == 0:
                raise ParsingError("The PDF document has no pages.")
            
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    
            text = "\n".join(text_parts).strip()
            if text:
                return text
        except ParsingError:
            raise
        except Exception as pypdf_err:
            raise ParsingError(
                f"PDF extraction failed. PyMuPDF failed to load ({str(e)}) "
                f"and pure-Python fallback also failed: {str(pypdf_err)}"
            )

    raise ParsingError("No text could be extracted from the PDF (it may be scanned/image-only).")

def parse_docx(file_bytes: bytes) -> str:
    """Parses a DOCX file in-memory using python-docx and returns its plain text."""
    if not file_bytes or len(file_bytes) == 0:
        raise ParsingError("The uploaded DOCX file is empty.")
    try:
        doc = docx.Document(BytesIO(file_bytes))
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables to ensure we don't miss content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            text_parts.append(p.text)
                            
        text = "\n".join(text_parts).strip()
        if not text:
            raise ParsingError("No text could be extracted from the DOCX document.")
        return text
    except ParsingError:
        raise
    except Exception as e:
        raise ParsingError(f"Corrupted or invalid DOCX document: {str(e)}")

def parse_resume(file_bytes: bytes, filename: str, content_type: str = "") -> str:
    """Dispatches parsing to appropriate parser based on file extension and content-type."""
    fn_lower = filename.lower()
    
    if fn_lower.endswith(".pdf") or content_type == "application/pdf":
        return parse_pdf(file_bytes)
    elif fn_lower.endswith(".docx") or content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword"
    ):
        return parse_docx(file_bytes)
    else:
        raise ParsingError("Unsupported file type. Only PDF and DOCX files are supported.")
